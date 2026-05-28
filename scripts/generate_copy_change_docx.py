# -*- coding: utf-8 -*-
"""
베베펫 6대 자산 카피 변경 결과를 워드 파일로 정리.
- 자산별 1페이지 분할
- 페이지 상단에 hover 상태 screenshot
- 본문은 2단(좌: 기존안, 우: 변경안) 테이블로 1:1 비교

출력: 바탕 화면\카피변경결과.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
SHOTS_DIR = os.path.join(ROOT, "docs", "screenshots")

DESKTOP = os.path.join(os.environ["USERPROFILE"], "OneDrive", "바탕 화면")
OUT_PATH = os.path.join(DESKTOP, "카피변경결과.docx")


# (제목, 기존안, 변경안) 트리플 묶음
ASSETS = [
    {
        "num": "01",
        "name_ko": "의료 자산",
        "name_en": "Medical",
        "screenshot": "card01_medical.png",
        "tagline_before": "단순한 분양을 넘어 생명을 책임지는 전문성을 상징합니다.",
        "tagline_after": "단순한 입양을 넘어, 생명을 끝까지 책임지는 메디컬 전문성입니다.",
        "points_before": [
            ("수의사 전담 관리", "입양 전 모든 아이의 건강검진을 자체 동물병원에서 수의사가 직접 실시합니다."),
            ("투명한 건강 정보", "검진 결과지 및 예방접종 기록을 투명하게 공개하여 입양 초기 불안감을 해소합니다."),
            ("생애주기 의료 지원", "입양 후에도 전문적인 의료 서비스를 지속적으로 제공받을 수 있는 든든한 파트너십을 제공합니다."),
        ],
        "points_after": [
            ("자체 의료진의 밀착 케어", "자체 메디컬 센터 수의사가 입양 전 모든 아이들을 직접 검진하고 보증합니다."),
            ("투명한 메디컬 데이터 공개", "검진 결과지와 예방접종 기록을 가감 없이 공개하여 초기 불안감을 해소합니다."),
            ("평생 안심 의료 파트너십", "입양에 그치지 않고 반려생활 전반의 전문 의료 서비스를 지속적으로 지원합니다."),
        ],
    },
    {
        "num": "02",
        "name_ko": "글로벌 자산",
        "name_en": "Global",
        "screenshot": "card02_global.png",
        "tagline_before": "유통의 한계를 넘어 최고의 제품을 합리적으로 공급합니다.",
        "tagline_after": "유통의 한계를 넘어, 세계 최고 수준의 프리미엄 제품을 가장 합리적으로 공급합니다.",
        "points_before": [
            ("전문가 소싱", "수의사가 성분을 분석하고 글로벌 사업부가 직접 계약한 해외 프리미엄 브랜드를 선보입니다."),
            ("신선도와 안전성", "중간 유통 단계를 생략하여 가장 신선한 사료를 가장 안전한 통관 절차를 거쳐 보호자에게 전달합니다."),
            ("독점적 가치", "국내에서 구하기 힘든 고품질 처방식이나 용품을 베베펫 고객만을 위해 독점 공급합니다."),
        ],
        "points_after": [
            ("수의사 검증 글로벌 소싱", "수의사가 성분을 분석하고 글로벌 사업부가 해외 본사와 직접 계약한 브랜드만 선보입니다."),
            ("다이렉트 유통 및 신선도", "중간 유통을 생략하고 안전한 통관을 거쳐 가장 신선한 상태로 보호자에게 전달합니다."),
            ("차별화된 독점 공급 가치", "국내에서 구하기 힘든 고품질 기능성 처방식과 프리미엄 용품을 독점 공급합니다."),
        ],
    },
    {
        "num": "03",
        "name_ko": "뷰티 자산",
        "name_en": "Beauty",
        "screenshot": "card03_beauty.png",
        "tagline_before": "반려동물의 위생과 스타일을 통해 삶의 질을 높입니다.",
        "tagline_after": "반려동물의 위생과 스타일을 넘어, 스트레스 없는 케어로 삶의 질을 높입니다.",
        "points_before": [
            ("전문 디자이너 케어", "견종·묘종별 특성을 이해하는 전문 미용사가 스트레스를 최소화한 미용 서비스를 제공합니다."),
            ("의료·미용 협업", "미용 중 피부나 건강 이상 발견 시 즉시 자체 의료진과 공유하는 '안전 미용 시스템'을 운영합니다."),
            ("", ""),  # 신규 항목 자리 (이전 항목 2개)
        ],
        "points_after": [
            ("펫 전용 하이엔드 그루밍", "품종별 특성을 깊이 이해하는 전문 디자이너가 스트레스를 최소화한 미용을 제공합니다."),
            ("메디컬 연계 안전 미용", "미용 중 피부나 관절 이상 발견 시 자체 의료진과 즉시 연계하여 대응합니다."),
            ("위생적인 도구·환경 관리 (신규)", "모든 미용 도구는 1회 사용 후 소독을 원칙으로 하며 위생적인 작업 공간을 유지합니다."),
        ],
        "points_note": "기존 2개 항목 유지 + 신규 1개 추가 (총 3개로 통일).",
    },
    {
        "num": "04",
        "name_ko": "시스템 자산",
        "name_en": "System",
        "screenshot": "card04_system.png",
        "tagline_before": "각 분야 전문가가 포진된 기업형 매니지먼트를 지향합니다.",
        "tagline_after": "개인이 아닌, 각 분야 전문가 그룹이 협업하는 체계적인 기업형 매니지먼트를 지향합니다.",
        "points_before": [
            ("분업화된 전문성", "경영지원, 글로벌 무역, 의료, 미용, 마케팅팀이 각자의 영역에서 최상의 결과물을 만듭니다."),
            ("표준화된 서비스", "어떤 채널을 통해 베베펫을 만나도 동일한 수준의 전문 상담과 서비스를 경험할 수 있는 매뉴얼 경영을 실천합니다."),
            ("", ""),  # 신규 항목 자리
        ],
        "points_after": [
            ("전문가 그룹의 협업 체계", "무역, 의료, 미용, 마케팅 등 고도화된 조직력으로 개인 숍과 차별화된 퀄리티를 만듭니다."),
            ("글로벌 스탠다드 매뉴얼", "어떤 채널에서 베베펫을 만나도 동일한 수준의 고품격 전문 서비스를 보장합니다."),
            ("체계적인 멤버십 매니지먼트 (신규)", "입양 이후에도 보호자와 반려동물의 생애 주기를 체계적으로 기록하고 관리하는 통합 시스템을 지향합니다."),
        ],
        "points_note": "기존 2개 항목 유지 + 신규 1개 추가 (총 3개로 통일).",
        "summary_before": "개인이 아닌, 각 분야의 / 전문가가 포진된 기업형 / 매니징을 지향합니다",
        "summary_after": "각 분야의 전문가가 / 포진된 기업형 매니징을 / 지향합니다",
        "summary_note": "카드 상단 3줄 요약에서 '개인이 아닌, ' 프리픽스 제거.",
    },
    {
        "num": "05",
        "name_ko": "신뢰 자산",
        "name_en": "Trust",
        "screenshot": "card05_trust.png",
        "tagline_before": "현장에서 쌓아온 실제 데이터와 고객과의 소통 기록입니다.",
        "tagline_after": "현장에서 증명된 실제 데이터와 보호자들과의 투명한 소통 기록입니다.",
        "points_before": [
            ("활발한 SNS 소통", "인스타그램 Reels를 통해 아이들의 노는 모습, 식사 시간, 매장 소독 현황 등을 매일 업데이트합니다."),
            ("실제 입양 후기", "'베베펫 가족'이 된 보호자들의 생생한 목소리와 건강하게 성장하는 아이들의 모습이 가장 강력한 신뢰의 증거입니다."),
            ("입양 데이터", "오픈 이후 쌓아온 수많은 입양 사례와 상담 기록은 지역사회에서 신뢰받는 브랜드로 자리 잡았음을 보여줍니다."),
        ],
        "points_after": [
            ("리얼타임 현장 투명 공개", "SNS를 통해 아이들의 일상 케어와 매장 위생 소독 현황을 매일 실시간으로 공유합니다."),
            ("보호자가 증명하는 가치", "실제 입양 가족들의 생생한 후기와 아이들의 성장 모습이 가장 확실한 신뢰의 증거입니다."),
            ("축적된 빅데이터의 힘", "오픈 이후 축적된 방대한 입양 및 상담 기록은 지역사회가 신뢰하는 기반입니다."),
        ],
    },
    {
        "num": "06",
        "name_ko": "철학 자산",
        "name_en": "Philosophy",
        "screenshot": "card06_philosophy.png",
        "tagline_before": "생명을 대하는 올바른 태도와 정직한 경영 원칙을 지향합니다.",
        "tagline_after": "생명 존중을 바탕으로 동물의 복지와 직원의 권리를 모두 존중하는 투명한 경영 원칙입니다.",
        "points_before": [
            ("독자적인 개체 관리", "쾌적하고 위생적인 매장 환경 유지 — 매일 정기 소독을 실시합니다."),
            ("연령별 맞춤 영양", "강아지·고양이의 생애 주기에 맞춘 필수 영양소 공급 및 전문가용 보충제 급여 원칙을 준수합니다."),
            ("전문 인력의 자부심", "수의사·미용사·무역 전문가가 체계적인 조직도 아래 책임감을 가지고 근무합니다."),
            ("노동법 준수", "투명한 노무 관리와 인센티브 구조로 직원이 행복하게 일할 환경을 만듭니다. 직원의 행복은 더 나은 케어로 이어집니다."),
        ],
        "points_after": [
            ("엄격한 위생·방역 원칙", "아이들이 머무는 모든 공간의 쾌적함을 위해 매일 정기 소독과 멸균 환경을 유지합니다."),
            ("생애주기별 맞춤 영양학", "월령과 건강 상태에 맞춘 필수 영양소 및 전문가용 보충제 급여 원칙을 철저히 준수합니다."),
            ("지속 가능한 상생 경영", "투명한 노무 관리로 전문가가 행복한 환경을 만듭니다. 직원의 행복은 최고 수준의 케어로 이어집니다."),
            ("", ""),  # 변경안은 3개
        ],
        "points_note": "이전 4개 항목 중 '전문 인력의 자부심' 삭제, '노동법 준수' → '지속 가능한 상생 경영'으로 라벨 교체. 4개 → 3개로 정리.",
    },
]


# ─── 서비스 카테고리 섹션 (ADOPT / HOSPITALITY / PRODUCT / BEAUTY / FOOD) ───
# 홈 #index 이하 풀-뷰포트 카테고리 프리뷰의 caption(헤드라인) + body(서브
# 카피) + 3개 points(label + 본문) 비교 데이터.
SECTIONS = [
    {
        "code": "ADOPT",
        "name_ko": "입양",
        "screenshot": "section_adopt.png",
        "caption_before": "건강한 첫만남",
        "caption_after": "완벽한 반려 라이프의 시작",
        "body_before": "전문 수의 검진을 거친 건강한 친구들과 평생을 함께할 책임 있는 분양.",
        "body_after": "전문 수의사 검진을 거친 건강한 가족을 안겨드리는 안심 입양",
        "points_before": [
            ("수의사 전담 관리", "자체 동물병원 수의사가 입양 전 건강검진을 직접 진행합니다"),
            ("투명한 건강 정보", "검진 결과와 예방접종 기록을 빠짐없이 투명하게 공개합니다"),
            ("생애주기 의료 지원", "입양 후에도 평생 동안 전문 의료 서비스를 책임지고 제공합니다"),
        ],
        "points_after": [
            ("라이프스타일 맞춤 상담", "보호자의 거주 환경과 일상을 종합적으로 고려하여 상담합니다"),
            ("검증된 브리더 협력", "생명 존중 철학에 공감하는 올바른 환경의 브리더와만 안전하게 협력합니다"),
            ("평생 안심 케어", "입양 후에도 평생 동안 전문 의료 서비스를 책임지고 제공합니다"),
        ],
        "points_note": "기존 3개 라벨이 '의료 자산'과 거의 동일해 한 화면 안 중복 — 입양 프로세스 자체(매칭·소싱·사후 케어)에 초점을 둔 3개로 전면 교체.",
    },
    {
        "code": "HOSPITALITY",
        "name_ko": "병원",
        "screenshot": "section_hospitality.png",
        "caption_before": "24시간 신뢰",
        "caption_after": "365일 · 24시간 안심 의료",
        "body_before": "전문 수의의 정기 진료부터 24시간 응급 케어까지, 신뢰의 의료 서비스.",
        "body_after": "정기 진료부터 24시간 응급 상황까지 믿을 수 있는 의료 서비스",
        "points_before": [
            ("의료팀 구성 공개", "진료·간호·원무팀이 각자 역할을 명확히 분담하여 운영합니다"),
            ("위생·소독 원칙", "쾌적한 환경을 위해 매일 정해진 시간에 매장 전체를 소독합니다"),
            ("예방접종·전문 수술 라인업", "예방접종부터 전문 수술까지 전 의료 서비스를 갖추고 있습니다"),
        ],
        "points_after": [
            ("전문 의료체계", "진료, 간호, 원무팀이 각자의 역할에 전문성을 더해 체계적으로 운영합니다"),
            ("철저한 원내 청정 방역", "쾌적하고 안전한 치유 환경을 위해 매일 정해진 시간에 원내 전체를 소독합니다"),
            ("종합 메디컬 솔루션", "필수 예방접종부터 난이도 높은 전문 수술까지 전방위 의료 서비스를 제공합니다"),
        ],
        "caption_note": "추상적 조어 '24시간 신뢰' → 구체적인 '365일 · 24시간 안심 의료'로 변경.",
    },
    {
        "code": "PRODUCT",
        "name_ko": "용품",
        "screenshot": "section_product.png",
        "caption_before": "일상의 도구",
        "caption_after": "건강한 일상을 위한 용품",
        "body_before": "케어, 외출, 휴식의 매 순간을 채우는 신중하게 고른 일상의 도구들.",
        "body_after": "매 순간을 안전하고 풍요롭게 채워줄 프리미엄 제품",
        "points_before": [
            ("카테고리별 직수입 라벨", "케어부터 잠자리까지 6개 전 카테고리를 직접 수입해 갖춥니다"),
            ("수의사 추천 마크", "모든 카테고리 용품을 수의사가 직접 확인하고 추천합니다"),
            ("통관 및 검역 안전성", "까다로운 통관과 검역을 모두 통과한 안전한 용품만 엄선합니다"),
        ],
        "points_after": [
            ("모든 일상을 위한 라인업", "케어, 외출, 휴식까지 반려생활에 필요한 모든 용품을 빠짐없이 한곳에 갖췄습니다"),
            ("엄격한 안심 기준 적용", "건강을 최우선으로 하여, 까다로운 유해 물질 및 성분 기준을 통과한 제품만 엄선합니다"),
            ("직접 사용해본 글로벌 직수입", "베베펫 수입팀이 직접 사용하고 검증한 해외 프리미엄 브랜드 제품만을 직수입하여 선보입니다"),
        ],
        "caption_note": "헤드라인에 '용품' 키워드 누락 → 영문 라벨 없이도 섹션 정체성이 드러나도록 보강.",
    },
    {
        "code": "BEAUTY & SPA",
        "name_ko": "미용·스파",
        "screenshot": "section_beauty.png",
        "caption_before": "교감의 시간",
        "caption_after": "전문 미용·스파",
        "body_before": "그루밍, 스파, 워시. 부드러운 손길 아래 깊어지는 교감의 시간.",
        "body_after": "견종과 묘종의 특성에 맞춘 안전하고 섬세한 전문 메디컬 케어",
        "points_before": [
            ("전문 디자이너 케어", "견종·묘종 특성을 이해한 미용사가 스트레스 없는 케어를 제공합니다"),
            ("의료·미용 협업", "미용 중 이상 발견 시 즉시 자체 의료진과 공유하여 대응합니다"),
            ("", ""),
        ],
        "points_after": [
            ("전문 디자이너 케어", "아이들의 특성을 완벽히 이해한 미용사가 스트레스 없는 1:1 맞춤 미용을 제공합니다"),
            ("의료·미용 협업 시스템", "미용 중 피부나 건강 이상 발견 시, 즉시 자체 의료진과 연계하여 신속하게 대응합니다"),
            ("교차 오염 없는 안심 위생 (신규)", "모든 도구의 철저한 멸균 소독 원칙을 준수하여 감염 위험을 차단합니다"),
        ],
        "points_note": "기존 2개 항목 유지 + '교차 오염 없는 안심 위생' 신규 추가 (다른 섹션과 같은 3개 구조로 통일).",
    },
    {
        "code": "FOOD",
        "name_ko": "사료·간식",
        "screenshot": "section_food.png",
        "caption_before": "엄선된 영양",
        "caption_after": "엄선된 영양 케어",
        "body_before": "프리미엄 사료부터 정성 가득 간식까지, 베베펫이 엄선한 안심의 영양.",
        "body_after": "프리미엄 사료부터 정성 가득한 간식까지, 깐깐하게 고른 푸드 셀렉션",
        "points_before": [
            ("전문가 소싱", "수의사가 성분을 분석한 해외 프리미엄 브랜드를 직접 들여옵니다"),
            ("독점적 가치", "국내에서 구하기 힘든 고품질 처방식을 독점적으로 공급합니다"),
            ("신선도와 안정성", "중간 유통을 생략하고 안전한 통관 절차로 신선하게 전달합니다"),
        ],
        "points_after": [
            ("수의사 성분 분석 소싱", "영양학적 기준과 원료의 안전성을 계량적으로 분석하여 우수제품을 선별합니다"),
            ("독점 공급 처방식 라인업", "시중에서 구하기 힘든 고품질 글로벌 처방식을 독점적으로 공급합니다"),
            ("입맛에 맞춘 기호성 테스트", "사전 급여 테스트를 거쳐 거부감 없이 맛있게 먹을 수 있는 제품만 엄선합니다"),
        ],
        "points_note": "'신선도와 안정성' 오타 정정 포함 (안정성 → 안전성 컨셉으로 라벨 자체 교체).",
    },
]


# ─── 헬퍼 ───────────────────────────────────────────────

def set_korean_font(run, name="맑은 고딕", size=10.5, bold=False, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc_or_cell, text, size=10.5, bold=False, color=None,
             align=None, space_before=0, space_after=0, line_spacing=None):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=bold, color=color)
    return p


def add_runs(paragraph, runs_spec):
    """[(text, kwargs), ...] 형태로 한 단락 안에 여러 run 추가."""
    for text, kwargs in runs_spec:
        r = paragraph.add_run(text)
        set_korean_font(r, **kwargs)


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def shade_cell(cell, hex_color):
    """셀 배경색 (예: 'F4F4F4')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, sides=("top", "bottom", "left", "right"),
                    color="DDDDDD", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for side in sides:
        elm = tc_borders.find(qn(f"w:{side}"))
        if elm is None:
            elm = OxmlElement(f"w:{side}")
            tc_borders.append(elm)
        elm.set(qn("w:val"), "single")
        elm.set(qn("w:sz"), size)
        elm.set(qn("w:color"), color)


def fill_compare_cell(cell, label_text, color_label, items, dim=False):
    """비교 테이블의 한 셀(기존/변경)을 채운다.
    items: [(label, body), ...]  label이 빈 문자열이면 '— 해당 항목 없음 —'.
    """
    # 셀 첫 단락에 라벨
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(4)
    add_runs(p0, [(label_text, dict(size=10, bold=True, color=color_label))])

    for lbl, body in items:
        if not lbl and not body:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_runs(p, [
                ("— 해당 항목 없음 —", dict(size=9.5, color=(0xBB, 0xBB, 0xBB))),
            ])
            continue
        # label 라인
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.25
        bullet_color = (0xAA, 0xAA, 0xAA) if dim else (0x10, 0x60, 0x4A)
        add_runs(p, [
            ("● ", dict(size=8, color=bullet_color)),
            (lbl, dict(size=10.5, bold=True,
                       color=(0x77, 0x77, 0x77) if dim else (0x10, 0x10, 0x10))),
        ])
        # body 라인
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.line_spacing = 1.3
        add_runs(p, [
            (body, dict(size=9.5,
                        color=(0x99, 0x99, 0x99) if dim else (0x33, 0x33, 0x33))),
        ])


def add_compare_block(doc, header_text, before, after, note=None,
                      structured=False):
    """라벨 + 2단 비교 테이블.
    structured=False면 before/after가 str. True면 [(label, body), ...] 리스트.
    """
    # 섹션 라벨
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_runs(p, [
        ("▎ ", dict(size=10, color=(0x10, 0x60, 0x4A))),
        (header_text, dict(size=11, bold=True, color=(0x22, 0x22, 0x22))),
    ])

    if note:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.4)
        add_runs(p, [
            ("※ ", dict(size=9, color=(0xC0, 0x85, 0x10))),
            (note, dict(size=9, color=(0x88, 0x55, 0x10))),
        ])

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # 동일 너비 두 컬럼
    col_w = Cm(8.2)
    for c in table.columns:
        c.width = col_w
    row = table.rows[0]
    for cell in row.cells:
        cell.width = col_w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_border(cell)

    left, right = row.cells
    # 좌측은 약간 연한 배경(기존안), 우측은 화이트 (변경안 — 강조)
    shade_cell(left, "F7F7F7")
    shade_cell(right, "FFFFFF")

    if structured:
        fill_compare_cell(left, "기존안 (BEFORE)", (0x99, 0x99, 0x99),
                          before, dim=True)
        fill_compare_cell(right, "변경안 (AFTER)", (0x10, 0x60, 0x4A),
                          after, dim=False)
    else:
        # 단일 문장 비교 (tagline, summary)
        fill_compare_cell(left, "기존안 (BEFORE)", (0x99, 0x99, 0x99),
                          [("", before)], dim=True)
        fill_compare_cell(right, "변경안 (AFTER)", (0x10, 0x60, 0x4A),
                          [("", after)], dim=False)


def add_screenshot(doc, image_path, width_cm=16):
    """페이지 폭에 맞춰 screenshot 삽입."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    add_runs(p, [
        (text, dict(size=9, color=(0xAA, 0xAA, 0xAA), name="맑은 고딕")),
    ])


# ─── 빌더 ───────────────────────────────────────────────

def build_doc():
    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ── 표지 ──
    add_para(doc, "BEBE PET", size=11, bold=True, color=(0x99, 0x99, 0x99),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "홈 화면 카피 변경 결과",
             size=26, bold=True, color=(0x10, 0x60, 0x4A),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc,
             "Part 1 — Hero stage 2 · 6대 자산  /  Part 2 — 서비스 카테고리 5종",
             size=11, color=(0x55, 0x55, 0x55),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "ADOPT · HOSPITALITY · PRODUCT · BEAUTY & SPA · FOOD",
             size=9, color=(0x99, 0x99, 0x99),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "2026.05.29 정리", size=10, color=(0xAA, 0xAA, 0xAA),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    # 개요 박스
    overview_table = doc.add_table(rows=1, cols=1)
    overview_table.autofit = False
    overview_table.columns[0].width = Cm(16.5)
    o_cell = overview_table.rows[0].cells[0]
    o_cell.width = Cm(16.5)
    shade_cell(o_cell, "F4F9F6")
    set_cell_border(o_cell, color="CFE2D6")

    p = o_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_runs(p, [
        ("개요 ", dict(size=11, bold=True, color=(0x10, 0x60, 0x4A))),
    ])
    p2 = o_cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.45
    add_runs(p2, [
        ("Part 1 — ", dict(size=10, bold=True, color=(0x10, 0x60, 0x4A))),
        ("Hero stage 2 6대 자산 카드 hover 상세의 ", dict(size=10, color=(0x33, 0x33, 0x33))),
        ("태그라인 + 세부 항목(label/body)", dict(size=10, bold=True, color=(0x33, 0x33, 0x33))),
        ("을 전면 개정. 자산별로 한 페이지에 hover 캡처와 기존/변경안 1:1 비교.",
         dict(size=10, color=(0x33, 0x33, 0x33))),
    ])
    p2b = o_cell.add_paragraph()
    p2b.paragraph_format.space_before = Pt(2)
    p2b.paragraph_format.line_spacing = 1.45
    add_runs(p2b, [
        ("Part 2 — ", dict(size=10, bold=True, color=(0x10, 0x60, 0x4A))),
        ("홈 #index 이하 서비스 카테고리 5종(", dict(size=10, color=(0x33, 0x33, 0x33))),
        ("ADOPT · HOSPITALITY · PRODUCT · BEAUTY & SPA · FOOD", dict(size=10, bold=True, color=(0x33, 0x33, 0x33))),
        (")의 헤드라인·서브카피·3개 포인트를 신규 톤으로 일괄 교체.", dict(size=10, color=(0x33, 0x33, 0x33))),
    ])
    p3 = o_cell.add_paragraph()
    p3.paragraph_format.space_before = Pt(6)
    p3.paragraph_format.line_spacing = 1.45
    add_runs(p3, [
        ("주요 정리 사항\n", dict(size=10, bold=True, color=(0x10, 0x60, 0x4A))),
        ("• 한국어 전 카피에서 '분양' → '입양'으로 용어 통일 (네비/SEO/본문/뉴스 일괄 적용)\n",
         dict(size=9.5, color=(0x55, 0x55, 0x55))),
        ("• 04 시스템 자산: 카드 상단 3줄 요약에서 '개인이 아닌, ' 프리픽스 제거\n",
         dict(size=9.5, color=(0x55, 0x55, 0x55))),
        ("• 03 뷰티 / 04 시스템: 6대 자산 항목 2개 → 3개로 통일 (각 1개 신규 추가)\n",
         dict(size=9.5, color=(0x55, 0x55, 0x55))),
        ("• 06 철학: 4개 → 3개로 정리 ('전문 인력의 자부심' 삭제, '노동법 준수' → '지속 가능한 상생 경영' 라벨 교체)\n",
         dict(size=9.5, color=(0x55, 0x55, 0x55))),
        ("• ADOPT: 의료 자산과 중복되던 3개 라벨 전면 교체 (입양 프로세스 중심으로 재구성)\n",
         dict(size=9.5, color=(0x55, 0x55, 0x55))),
        ("• HOSPITALITY: '24시간 신뢰' → '365일 · 24시간 안심 의료' (구체화)\n",
         dict(size=9.5, color=(0x55, 0x55, 0x55))),
        ("• PRODUCT: '일상의 도구' → '건강한 일상을 위한 용품' (용품 키워드 보강)\n",
         dict(size=9.5, color=(0x55, 0x55, 0x55))),
        ("• BEAUTY: '교감의 시간' → '전문 미용·스파' + 항목 2개 → 3개\n",
         dict(size=9.5, color=(0x55, 0x55, 0x55))),
        ("• FOOD: '안정성' 오타 정정 포함 라벨 전면 교체",
         dict(size=9.5, color=(0x55, 0x55, 0x55))),
    ])

    # ── Part 1 헤더 ──
    page_break(doc)
    add_para(doc, "PART 1", size=11, bold=True, color=(0x10, 0x60, 0x4A),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "Hero stage 2 · 베베펫 6대 자산",
             size=22, bold=True, color=(0x22, 0x22, 0x22),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "카드 hover 시 상단에 노출되는 상세 카피 — 자산별 1페이지",
             size=10, color=(0x88, 0x88, 0x88),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # 자산별 페이지
    for asset in ASSETS:
        page_break(doc)

        # 헤더: 번호 · EN · KR
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_runs(p, [
            (asset["num"], dict(size=22, bold=True, color=(0x10, 0x60, 0x4A))),
            ("  ·  ", dict(size=14, color=(0xCC, 0xCC, 0xCC))),
            (asset["name_en"], dict(size=16, bold=True, color=(0x44, 0x44, 0x44))),
            ("   ", dict(size=14)),
            (asset["name_ko"], dict(size=16, color=(0x77, 0x77, 0x77))),
        ])
        # 구분선
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_after = Pt(6)
        add_runs(p_line, [
            ("─" * 70, dict(size=8, color=(0xE0, 0xE0, 0xE0))),
        ])

        # 스크린샷
        shot_path = os.path.join(SHOTS_DIR, asset["screenshot"])
        if os.path.exists(shot_path):
            add_screenshot(doc, shot_path, width_cm=16)
            add_caption(doc,
                        f"▲ 홈 → Hero stage 2 → {asset['num']} · {asset['name_ko']} 카드 hover 상태")
        else:
            add_para(doc, f"(screenshot missing: {asset['screenshot']})",
                     size=9, color=(0xCC, 0x55, 0x55),
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

        # 태그라인 비교 (단일 문장)
        add_compare_block(doc, "태그라인 (Hover 시 큰 글씨)",
                          asset["tagline_before"], asset["tagline_after"])

        # summary 변경이 있으면 추가 (04 시스템만 해당)
        if "summary_before" in asset:
            add_compare_block(doc, "카드 상단 3줄 요약",
                              asset["summary_before"], asset["summary_after"],
                              note=asset.get("summary_note"))

        # 세부 points 비교 (3행)
        add_compare_block(doc, "세부 항목 (label + 본문)",
                          asset["points_before"], asset["points_after"],
                          note=asset.get("points_note"),
                          structured=True)

    # ── Part 2 헤더 ──
    page_break(doc)
    add_para(doc, "PART 2", size=11, bold=True, color=(0x10, 0x60, 0x4A),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "홈 #index 서비스 카테고리 섹션",
             size=22, bold=True, color=(0x22, 0x22, 0x22),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "ADOPT · HOSPITALITY · PRODUCT · BEAUTY & SPA · FOOD — 섹션별 1페이지",
             size=10, color=(0x88, 0x88, 0x88),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # 카테고리 섹션별 페이지
    for sec in SECTIONS:
        page_break(doc)

        # 헤더: EN code + KR name
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_runs(p, [
            (sec["code"], dict(size=18, bold=True, color=(0x10, 0x60, 0x4A))),
            ("   ", dict(size=14)),
            (sec["name_ko"], dict(size=16, color=(0x77, 0x77, 0x77))),
        ])
        # 구분선
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_after = Pt(6)
        add_runs(p_line, [
            ("─" * 70, dict(size=8, color=(0xE0, 0xE0, 0xE0))),
        ])

        # 스크린샷
        shot_path = os.path.join(SHOTS_DIR, sec["screenshot"])
        if os.path.exists(shot_path):
            add_screenshot(doc, shot_path, width_cm=16)
            add_caption(doc,
                        f"▲ 홈 → #{sec['code'].split()[0].lower().replace('hospitality','location').replace('food','partners')} 섹션 (실제 노출 화면)")
        else:
            add_para(doc, f"(screenshot missing: {sec['screenshot']})",
                     size=9, color=(0xCC, 0x55, 0x55),
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

        # caption(헤드라인) 비교
        add_compare_block(doc, "헤드라인 (caption)",
                          sec["caption_before"], sec["caption_after"],
                          note=sec.get("caption_note"))

        # body(서브카피) 비교
        add_compare_block(doc, "서브카피 (body)",
                          sec["body_before"], sec["body_after"])

        # 3개 points 비교
        add_compare_block(doc, "세부 항목 (label + 본문)",
                          sec["points_before"], sec["points_after"],
                          note=sec.get("points_note"),
                          structured=True)

    # 푸터
    page_break(doc)
    add_para(doc, "─" * 30, size=9, color=(0xCC, 0xCC, 0xCC),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=18, space_after=4)
    add_para(doc, "본 문서는 코드 반영 완료된 카피를 기준으로 작성되었습니다.",
             size=9, color=(0x99, 0x99, 0x99),
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, f"스크린샷: {SHOTS_DIR}",
             size=8, color=(0xBB, 0xBB, 0xBB),
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(f"OK: {path}")
